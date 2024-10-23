import os

import joblib as joblib
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR
from sklearn.neighbors import KNeighborsRegressor
from sklearn.preprocessing import LabelEncoder
import aspose.words as aw

# Wczytanie danych
url = 'https://vincentarelbundock.github.io/Rdatasets/csv/AER/CollegeDistance.csv'
data = pd.read_csv(url)
nazwaWord='analiza_statystyczna.docx'
nazwaPDF=nazwaWord

# Oczyszczanie starych plikow z zeszłych uruchomien jesli takie były
if os.path.exists(nazwaWord):
    os.remove(nazwaWord)

# Utworzenie dokumentu Word
doc = Document()
doc.add_heading('Analiza statystyczna zmiennych', 0)

# Wstęp
doc.add_heading('Wstęp', level=1)
doc.add_paragraph('Niniejszy dokument przedstawia wyniki analizy statystycznej danych, które pochodzą z zestawu CollegeDistance. '
                  'Dane te zostały wstępnie przetworzone, a następnie posłużyły do budowy modelu predykcyjnego, którego celem jest przewidywanie zmiennej "score". '
                  'W trakcie pracy nad projektem przeprowadzono oczyszczanie danych, analizę statystyczną, wybór modeli, ich ocenę oraz optymalizację.'
                  'Od raz chcę ostrzec że po wczesniejszych badaniach najlepszy wyniki osiągneła regresja liniowa i pod to głównie został przygotowany plik optymalizujacy model')

doc.add_heading('Informacje o zmiennyh', level=1)

# Funkcja tworząca wykresy i statystyki dla każdej zmiennej
def analyze_column(column_name):
    doc.add_heading(f'Zmienna: {column_name}', level=1)
    doc.add_paragraph(f'Typ zmiennej: {data[column_name].dtype}')

    # Statystyki opisowe
    desc_stats = data[column_name].describe()
    doc.add_paragraph(f'Statystyki opisowe:\n{desc_stats.to_string()}')

    # Tworzenie wykresu histogramu dla zmiennej
    plt.figure(figsize=(10, 6))
    sns.histplot(data[column_name], bins=30, kde=True)
    plt.title(f'Rozkład zmiennej {column_name}')
    plt.xlabel(column_name)
    plt.ylabel('Liczba wystąpień')

    # Zapis wykresu do pliku tymczasowego
    tmp_filename = f'{column_name}_distribution.png'
    plt.savefig(tmp_filename)
    plt.close()

    # Dodanie wykresu do dokumentu
    doc.add_paragraph(f'Wykres histogramu dla zmiennej {column_name}:')
    doc.add_picture(tmp_filename, width=Inches(5.0))

    # Usunięcie pliku tymczasowego
    os.remove(tmp_filename)


# Sprawdzenie dostępnych kolumn i wykonanie analizy dla każdej zmiennej
for col in data.columns:
    analyze_column(col)


# Zapisanie dokumentu przed oczyszczaniem danych
doc.save(nazwaWord)

# Usunięcie kolumny 'rownames'
data.drop('rownames', axis=1, inplace=True)

# Zakodowanie kolumny 'income' na wartości liczbowe
label_encoder = LabelEncoder()
data['income'] = label_encoder.fit_transform(data['income'])

# Zakodowanie kolumny 'ethnicity' na wartości liczbowe
label_encoder = LabelEncoder()
data['ethnicity'] = label_encoder.fit_transform(data['ethnicity'])

# Zamiana wartości 'male' i 'female' na 1 i 0 w kolumnie 'gender'
data['gender'] = data['gender'].map({'male': 1, 'female': 0})

# Zamiana wartości 'yes' i 'no' na 1 i 0 w odpowiednich kolumnach
columns_to_convert = ['fcollege', 'mcollege', 'home', 'urban']
for col in columns_to_convert:
    data[col] = data[col].map({'yes': 1, 'no': 0})

# Zamiana wartości 'region' na 1 dla 'other' i 0 dla 'high', 'low'
data['region'] = data['region'].map({'other': 1, 'west': 0})

# Sprawdzenie struktury danych
print("\nInformacje o danych:")
print(data.info())

# Sprawdzenie brakujących wartości
print("\nBrakujące wartości w danych:")
print(data.isnull().sum())

# Obsługa brakujących wartości - usunięcie wierszy z więcej niż 3 brakującymi wartościami
initial_shape = data.shape
data.dropna(thresh=len(data.columns) - 3, inplace=True)  # Usunięcie wierszy z więcej niż 3 brakującymi wartościami
print(f"\nKształt danych przed usunięciem wierszy: {initial_shape}")
print(f"Kształt danych po usunięciu wierszy: {data.shape}")

# Uzupełnienie brakujących wartości medianą dla kolumn numerycznych
numerical_cols = data.select_dtypes(include=['float64', 'int64']).columns.tolist()
for col in numerical_cols:
    median_value = data[col].median()
    data[col] = data[col].fillna(median_value)
    print(f"Uzupełniono brakujące wartości w kolumnie '{col}' medianą: {median_value}")

# Sprawdzenie, czy są nadal brakujące wartości
print("\nBrakujące wartości w danych po imputacji:")
print(data.isnull().sum())

# Zapisanie oczyszczonych danych
data.to_csv('cleaned_data.csv', index=False)
print("\nDane oczyszczone i zapisane do 'cleaned_data.csv'.")

# Etap porównania modeli
doc.add_heading('Porównanie modeli', level=1)

# Wybór zmiennych do modelowania
X = data.drop(columns=['score'])
y = data['score']

# Podział danych na zbiór treningowy i testowy
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Lista modeli do porównania
models = {
    'Linear Regression': LinearRegression(),
    'Decision Tree': DecisionTreeRegressor(),
    'Random Forest': RandomForestRegressor(),
    'Support Vector Regressor': SVR(),
    'K-Neighbors Regressor': KNeighborsRegressor()
}

best_model_name = None
best_mse = float('inf')
model_results = []

# Porównanie modeli
for model_name, model in models.items():
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    mse = mean_squared_error(y_test, y_pred)
    model_results.append((model_name, mse))
    doc.add_paragraph(f'Model: {model_name}, Mean Squared Error (MSE): {mse:.4f}')

    if mse < best_mse:
        best_mse = mse
        best_model_name = model_name
        best_model = model

# Zapisanie najlepszego modelu
joblib.dump(best_model, 'best_model.pkl')
doc.add_paragraph(f'Najlepszy model: {best_model_name} z MSE: {best_mse:.4f}')

doc.add_heading('Podsumowanie', level=1)
doc.add_paragraph(f'W niniejszej pracy wykonano pełną analizę danych, w tym ich oczyszczenie i przygotowanie do modelowania. '
                  f'Zastosowano kilka modeli predykcyjnych, spośród których najlepszy okazał się model "{best_model_name}" z najniższym błędem średniokwadratowym (MSE): {best_mse:.4f}. '
                  'Dokument ten zawiera zarówno szczegóły dotyczące poszczególnych zmiennych, jak i oceny porównywanych modeli. '
                  'W razie potrzeby można przeprowadzić dodatkową optymalizację, aby jeszcze bardziej poprawić jakość modelu.')
doc.save(nazwaWord)
print(f'Dokument zapisany jako {nazwaWord}.')
doc_pdf = aw.Document(nazwaWord)
doc_pdf.save('analiza_statystyczna.pdf')
print('zapsano pdf')
