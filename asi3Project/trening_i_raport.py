import os

import joblib
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, GridSearchCV, cross_val_score
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
from sklearn.linear_model import Ridge
from docx import Document

# Wczytanie oczyszczonych danych
data = pd.read_csv('cleaned_data.csv')

# Wybór zmiennych do modelowania
X = data.drop(columns=['score'])
y = data['score']

# Podział danych na zbiór treningowy i testowy
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Wczytanie najlepszego modelu
try:
    best_model = joblib.load('best_model.pkl')
except FileNotFoundError:
    raise FileNotFoundError("Plik z modelem nie został znaleziony.")

# Ocena modelu
y_pred = best_model.predict(X_test)

# Obliczenie metryk
mse = mean_squared_error(y_test, y_pred)
mae = mean_absolute_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)

# Wyswietlanie wyników w konsoli
print(f'Ocena modelu: \nMSE: {mse:.4f}, MAE: {mae:.4f}, R²: {r2:.4f}')

# Tworzenie dokumentu Word
doc = Document()
doc.add_heading('Optymalizacja Modelu', 0)

# Wstęp do dokumentu
doc.add_paragraph('W tej części dokonano oceny oraz optymalizacji modelu predykcyjnego.')

# Wstawienie wyników do dokumentu
doc.add_heading('Ocena modelu', level=1)
doc.add_paragraph(f'MSE: {mse:.4f}')
doc.add_paragraph(f'MAE: {mae:.4f}')
doc.add_paragraph(f'R²: {r2:.4f}')

# Optymalizacja (tuning hiperparametrów) jeśli wyniki nie są satysfakcjonujące
if mse > 45:
    print("Wyniki nie są satysfakcjonujące. Rozpoczynam optymalizację modelu.")

    # Optymalizacja Ridge Regression
    param_grid = {
        'alpha': [0.1, 1.0, 10.0]
    }

    grid_search = GridSearchCV(estimator=Ridge(), param_grid=param_grid,
                               scoring='neg_mean_squared_error', cv=5)

    grid_search.fit(X_train, y_train)

    # Walidacja krzyżowa
    cv_scores = cross_val_score(grid_search.best_estimator_, X_train, y_train, cv=5, scoring='neg_mean_squared_error')
    mean_cv_mse = -cv_scores.mean()

    # Wyswietlanie wyników optymalizacji w konsoli
    print(f'Walidacja krzyżowa - MSE: {mean_cv_mse:.4f} +/- {cv_scores.std():.4f}')
    print(f'Najlepsze parametry: {grid_search.best_params_}')

    best_params = grid_search.best_params_
    optimized_model = grid_search.best_estimator_

    # Ocena zoptymalizowanego modelu
    y_pred_optimized = optimized_model.predict(X_test)

    # Obliczenie metryk dla zoptymalizowanego modelu
    optimized_mse = mean_squared_error(y_test, y_pred_optimized)
    optimized_mae = mean_absolute_error(y_test, y_pred_optimized)
    optimized_r2 = r2_score(y_test, y_pred_optimized)

    # Wyswietlanie zoptymalizowanych wyników w konsoli
    print(f'Zoptymalizowany model: \nMSE: {optimized_mse:.4f}, MAE: {optimized_mae:.4f}, R²: {optimized_r2:.4f}')

    # Wstawienie zoptymalizowanych wyników do dokumentu
    doc.add_heading('Optymalizacja modelu', level=1)
    doc.add_paragraph(f'MSE (zoptymalizowany): {optimized_mse:.4f}')
    doc.add_paragraph(f'MAE (zoptymalizowany): {optimized_mae:.4f}')
    doc.add_paragraph(f'R² (zoptymalizowany): {optimized_r2:.4f}')
    doc.add_paragraph(f'Najlepsze parametry: {best_params}')

else:
    print("Model spełnia wymagania jakościowe. Optymalizacja nie jest konieczna.")
    doc.add_paragraph("Model spełnia wymagania jakościowe. Optymalizacja nie była konieczna.")

# Zapisanie dokumentu
doc.save('optymalizacja_modelu.docx')
print("Files in the current directory:", os.listdir('.'))
